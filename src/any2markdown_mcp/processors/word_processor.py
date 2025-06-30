"""
Word文档处理器

基于python-docx实现Word转换功能，支持：
- Word文档转Markdown转换
- 页眉页脚智能移除
- 图片提取和静态化
- 表格转换
- 样式保持
"""

import asyncio
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import re
import zipfile
import xml.etree.ElementTree as ET

import structlog
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls, qn
from PIL import Image
import io

from .base_processor import BaseProcessor

logger = structlog.get_logger(__name__)


class WordProcessor(BaseProcessor):
    """Word文档处理器"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        
        # Word特定配置
        self.preserve_formatting = self._get_config_value("word_preserve_formatting", True)
        self.convert_equations = self._get_config_value("word_convert_equations", True)
        
        logger.info("WordProcessor initialized")
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return ["docx", "doc"]
    
    async def convert(self, file_content: bytes, options: Dict[str, Any]) -> Dict[str, Any]:
        """
        转换Word文档
        
        Args:
            file_content: Word文件内容
            options: 转换选项
            
        Returns:
            转换结果
        """
        temp_files = []
        
        try:
            # 验证文件大小
            self.validate_file_size(file_content)
            
            # 保存临时Word文件
            temp_word_path = await self.save_temp_file(file_content, suffix=".docx")
            temp_files.append(str(temp_word_path))
            
            logger.info("Starting Word conversion", 
                       file_size=len(file_content),
                       temp_path=str(temp_word_path))
            
            # 打开Word文档
            doc = Document(str(temp_word_path))
            
            # 分析文档结构
            doc_info = await self._analyze_document_structure(doc)
            
            # 提取图片（如果启用）
            images_info = []
            if options.get("extract_images", True):
                images_info = await self._extract_images(temp_word_path)
                # 注意：不要将图片文件添加到temp_files中，因为它们需要被保留用于静态文件服务
            
            # 转换文档内容
            markdown_content = await self._convert_document_content(doc, options)
            
            # 移除页眉页脚（如果启用）
            if options.get("remove_header_footer", True):
                markdown_content = await self._remove_header_footer(
                    markdown_content, doc_info, options
                )
            
            # 嵌入图片URL到markdown中
            if images_info:
                markdown_content = await self._embed_image_urls(markdown_content, images_info)
            
            # 分页处理（如果启用）
            if options.get("paginate_output", True):
                pages = await self._paginate_content(markdown_content, doc_info)
            else:
                pages = [{"page": 1, "content": markdown_content}]
            
            # 构建元数据
            metadata = {
                "source_type": "word",
                "paragraph_count": doc_info["paragraph_count"],
                "table_count": doc_info["table_count"],
                "images": images_info,
                "file_size": len(file_content),
                "options_used": options,
                "document_properties": doc_info.get("properties", {})
            }
            
            # 格式化输出
            result = self.format_output(
                markdown_content if not options.get("paginate_output", True) else pages,
                metadata,
                options.get("output_format", "markdown")
            )
            
            logger.info("Word conversion completed successfully",
                       paragraphs_processed=doc_info["paragraph_count"],
                       images_extracted=len(images_info))
            
            return result
            
        except Exception as e:
            logger.error("Word conversion failed", error=str(e))
            raise
        finally:
            # 清理临时文件
            if temp_files:
                await self.cleanup_temp_files(temp_files)
    
    async def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        """分析Word文档结构"""
        try:
            info = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "has_images": False,
                "has_headers_footers": False,
                "styles_used": set(),
                "properties": {}
            }
            
            # 文档属性
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                info["properties"]["title"] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                info["properties"]["author"] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                info["properties"]["created"] = str(doc.core_properties.created)
            
            # 检查样式使用情况
            for paragraph in doc.paragraphs:
                if paragraph.style and paragraph.style.name:
                    info["styles_used"].add(paragraph.style.name)
            
            # 检查是否有图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    info["has_images"] = True
                    break
            
            # 检查页眉页脚
            for section in doc.sections:
                if (section.header.paragraphs and 
                    any(p.text.strip() for p in section.header.paragraphs)):
                    info["has_headers_footers"] = True
                    break
                if (section.footer.paragraphs and 
                    any(p.text.strip() for p in section.footer.paragraphs)):
                    info["has_headers_footers"] = True
                    break
            
            info["styles_used"] = list(info["styles_used"])
            
            logger.debug("Word document structure analyzed", info=info)
            return info
            
        except Exception as e:
            logger.error("Failed to analyze Word document structure", error=str(e))
            raise
    
    async def _extract_images(self, word_path: Path) -> List[Dict[str, Any]]:
        """从Word文档中提取图片"""
        images_info = []
        
        try:
            # Word文档实际上是一个ZIP文件
            with zipfile.ZipFile(str(word_path), 'r') as docx_zip:
                # 查找图片文件
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             any(f.lower().endswith(ext) for ext in 
                                 ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'])]
                
                for img_index, img_path in enumerate(image_files):
                    try:
                        # 读取图片数据
                        img_data = docx_zip.read(img_path)
                        
                        # 获取原始文件名和扩展名
                        original_filename = Path(img_path).name
                        
                        # 保存图片（使用实例ID避免文件名冲突）
                        filename = f"word_{self.instance_id}_img_{img_index + 1}_{original_filename}"
                        image_info = await self.save_image(img_data, filename)
                        
                        # 添加额外信息
                        image_info.update({
                            "index": img_index + 1,
                            "original_path": img_path,
                            "original_filename": original_filename
                        })
                        
                        images_info.append(image_info)
                        
                        logger.debug("Image extracted from Word", 
                                   index=img_index + 1,
                                   filename=filename)
                        
                    except Exception as e:
                        logger.warning("Failed to extract image from Word", 
                                     path=img_path,
                                     error=str(e))
                        continue
            
            logger.info("Image extraction from Word completed", count=len(images_info))
            return images_info
            
        except Exception as e:
            logger.error("Failed to extract images from Word document", error=str(e))
            return []
    
    async def _convert_document_content(self, doc: Document, options: Dict[str, Any]) -> str:
        """转换Word文档内容为Markdown"""
        try:
            markdown_lines = []
            
            # 处理文档标题（如果有）
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                markdown_lines.append(f"# {doc.core_properties.title}")
                markdown_lines.append("")
            
            # 处理段落和表格
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    paragraph = self._find_paragraph_by_element(doc, element)
                    if paragraph:
                        markdown_line = await self._convert_paragraph(paragraph, options)
                        if markdown_line:
                            markdown_lines.append(markdown_line)
                
                elif element.tag.endswith('tbl'):  # 表格
                    table = self._find_table_by_element(doc, element)
                    if table:
                        table_markdown = await self._convert_table(table, options)
                        if table_markdown:
                            markdown_lines.extend(table_markdown)
                            markdown_lines.append("")  # 表格后添加空行
            
            return "\n".join(markdown_lines)
            
        except Exception as e:
            logger.error("Failed to convert Word document content", error=str(e))
            raise
    
    def _find_paragraph_by_element(self, doc: Document, element) -> Optional[object]:
        """根据XML元素查找对应的段落对象"""
        for paragraph in doc.paragraphs:
            if paragraph._element == element:
                return paragraph
        return None
    
    def _find_table_by_element(self, doc: Document, element) -> Optional[object]:
        """根据XML元素查找对应的表格对象"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None
    
    async def _convert_paragraph(self, paragraph, options: Dict[str, Any]) -> str:
        """转换段落为Markdown"""
        if not paragraph.text.strip():
            return ""
        
        text = paragraph.text.strip()
        
        # 根据样式转换
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        # 标题样式
        if "Heading" in style_name:
            level = 1
            if "Heading 1" in style_name:
                level = 1
            elif "Heading 2" in style_name:
                level = 2
            elif "Heading 3" in style_name:
                level = 3
            elif "Heading 4" in style_name:
                level = 4
            elif "Heading 5" in style_name:
                level = 5
            elif "Heading 6" in style_name:
                level = 6
            
            return f"{'#' * level} {text}"
        
        # 列表项
        elif "List" in style_name or paragraph.style.name in ["List Paragraph"]:
            return f"- {text}"
        
        # 引用
        elif "Quote" in style_name:
            return f"> {text}"
        
        # 代码
        elif "Code" in style_name:
            return f"```\n{text}\n```"
        
        # 普通段落
        else:
            # 处理内联格式
            formatted_text = await self._process_inline_formatting(paragraph, options)
            return formatted_text
    
    async def _process_inline_formatting(self, paragraph, options: Dict[str, Any]) -> str:
        """处理段落内的内联格式"""
        if not options.get("preserve_formatting", True):
            return paragraph.text
        
        result = ""
        
        for run in paragraph.runs:
            text = run.text
            
            # 应用格式
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            
            result += text
        
        return result
    
    async def _convert_table(self, table, options: Dict[str, Any]) -> List[str]:
        """转换表格为Markdown"""
        if not table.rows:
            return []
        
        markdown_lines = []
        
        # 处理表头
        header_row = table.rows[0]
        headers = []
        for cell in header_row.cells:
            headers.append(cell.text.strip() or " ")
        
        markdown_lines.append("| " + " | ".join(headers) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 处理数据行
        for row in table.rows[1:]:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                cells.append(cell_text or " ")
            
            # 确保列数一致
            while len(cells) < len(headers):
                cells.append(" ")
            
            markdown_lines.append("| " + " | ".join(cells[:len(headers)]) + " |")
        
        return markdown_lines
    
    async def _remove_header_footer(self, content: str, doc_info: Dict[str, Any], options: Dict[str, Any]) -> str:
        """移除Word文档的页眉页脚"""
        if not options.get("remove_header_footer", True):
            return content
        
        logger.info("Removing header/footer from Word content")
        
        # 使用基类的通用方法
        cleaned_content = self.remove_header_footer_text(content, options)
        
        # Word特定的页眉页脚处理
        cleaned_content = await self._remove_word_specific_headers_footers(cleaned_content, doc_info)
        
        return cleaned_content
    
    async def _remove_word_specific_headers_footers(self, content: str, doc_info: Dict[str, Any]) -> str:
        """移除Word特定的页眉页脚"""
        lines = content.split('\n')
        cleaned_lines = []
        
        # Word文档中的页眉页脚通常包含特定的格式模式
        word_header_footer_patterns = [
            r'^.*第\s*\d+\s*页.*共\s*\d+\s*页.*$',  # 中文页码格式
            r'^.*Page\s+\d+\s+of\s+\d+.*$',  # 英文页码格式
            r'^.*\d+\s*/\s*\d+.*$',  # 简单页码格式
            r'^.*\d{4}年\d{1,2}月\d{1,2}日.*$',  # 中文日期格式
            r'^.*\d{1,2}/\d{1,2}/\d{4}.*$',  # 英文日期格式
        ]
        
        for line in lines:
            line_stripped = line.strip()
            
            # 检查是否匹配页眉页脚模式
            is_header_footer = False
            for pattern in word_header_footer_patterns:
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    is_header_footer = True
                    logger.debug("Removing Word header/footer", line=line_stripped)
                    break
            
            if not is_header_footer:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    async def _embed_image_urls(self, content: str, images_info: List[Dict[str, Any]]) -> str:
        """将图片URL嵌入到markdown内容中"""
        if not images_info:
            return content
        
        logger.info("Embedding image URLs in Word markdown", count=len(images_info))
        
        # 在文档末尾添加图片
        lines = content.split('\n')
        
        # 添加图片部分
        lines.append('')
        lines.append('## Images')
        lines.append('')
        
        for img in images_info:
            alt_text = f"Image {img['index']}: {img.get('original_filename', 'image')}"
            lines.append(f"![{alt_text}]({img['url']})")
            lines.append('')
        
        return '\n'.join(lines)
    
    async def _paginate_content(self, content: str, doc_info: Dict[str, Any]) -> List[Dict[str, Any]]:
        """将内容分页（Word文档通常按段落或章节分页）"""
        pages = []
        
        # 按标题分页
        lines = content.split('\n')
        current_page_content = []
        current_page_num = 1
        
        for line in lines:
            # 检测主要标题作为分页点
            if line.startswith('# ') and current_page_content:
                # 保存当前页面
                pages.append({
                    "page": current_page_num,
                    "content": '\n'.join(current_page_content).strip()
                })
                
                # 开始新页面
                current_page_num += 1
                current_page_content = [line]
            else:
                current_page_content.append(line)
        
        # 保存最后一页
        if current_page_content:
            pages.append({
                "page": current_page_num,
                "content": '\n'.join(current_page_content).strip()
            })
        
        # 如果没有找到标题分隔符，将整个内容作为一页
        if not pages:
            pages.append({
                "page": 1,
                "content": content.strip()
            })
        
        logger.info("Word content paginated", total_pages=len(pages))
        return pages 